"""DOCX import and export using an explicit, portable document convention."""

from __future__ import annotations

import hashlib
import mimetypes
from pathlib import Path
from typing import Any

from ..models import Answer, Asset, Choice, Content, ContentBlock, Question, QuestionSet


def _require_docx():
    try:
        import docx
    except ImportError as exc:
        raise RuntimeError(
            "DOCX support requires: pip install 'dq-questionbank-core[docx]'"
        ) from exc
    return docx


def _append_content(
    paragraph, content: Content, assets: dict[str, Asset], assets_base: Path
) -> None:
    for block in content.blocks:
        if block.type == "text":
            paragraph.add_run(block.text or "")
        elif block.type == "math":
            paragraph.add_run(f"${block.latex or ''}$")
        elif block.type == "code":
            paragraph.add_run(block.text or "")
        elif block.type == "line_break":
            paragraph.add_run().add_break()
        elif block.type == "image":
            asset = assets.get(block.asset_id or "")
            if asset and not Path(asset.uri).is_absolute() and "://" not in asset.uri:
                image_path = (assets_base / asset.uri).resolve()
                try:
                    image_path.relative_to(assets_base.resolve())
                    if image_path.is_file():
                        paragraph.add_run(f"[Image: {block.alt_text or asset.id}] ")
                        paragraph.add_run().add_picture(str(image_path))
                        continue
                except (OSError, ValueError):
                    pass
            paragraph.add_run(f"[Image asset: {block.asset_id or 'missing'}]")
        elif block.type == "table":
            paragraph.add_run(" | ".join(" / ".join(row) for row in (block.rows or [])))


class DocxExporter:
    format_name = "docx"
    extensions = (".docx",)

    def dump(self, question_set: QuestionSet, target: Path, **options: Any) -> None:
        docx = _require_docx()
        document = docx.Document()
        document.core_properties.title = question_set.title
        document.core_properties.subject = "Structured educational questions"
        document.add_heading(question_set.title, level=0)
        if question_set.description:
            document.add_paragraph(question_set.description)
        assets_base = Path(options.get("assets_base") or target.parent)
        for question in question_set.questions:
            self._write_question(document, question, assets_base)
        target.parent.mkdir(parents=True, exist_ok=True)
        document.save(target)

    def _write_question(
        self,
        document,
        question: Question,
        assets_base: Path,
        *,
        level: int = 1,
        parent_id: str | None = None,
    ) -> None:
        document.add_heading(f"Question {question.id}", level=min(level, 9))
        marker_text = f"Type: {question.type} | Language: {question.language}"
        if parent_id:
            marker_text += f" | Parent: {parent_id}"
        document.add_paragraph(marker_text)
        assets = {item.id: item for item in question.assets}
        stem = document.add_paragraph("Question: ")
        _append_content(stem, question.stem, assets, assets_base)
        for choice in question.choices:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(f"[{choice.id}] ")
            _append_content(paragraph, choice.content, assets, assets_base)
        if question.answer:
            document.add_paragraph(f"Answer: {question.answer.kind}:{question.answer.value}")
        if question.solution:
            paragraph = document.add_paragraph("Solution: ")
            _append_content(paragraph, question.solution, assets, assets_base)
        for hint in question.hints:
            paragraph = document.add_paragraph("Hint: ")
            _append_content(paragraph, hint, assets, assets_base)
        for subquestion in question.subquestions:
            self._write_question(
                document, subquestion, assets_base, level=level + 1, parent_id=question.id
            )


class DocxImporter:
    format_name = "docx"
    extensions = (".docx",)

    def load(self, source: Path, **options: Any) -> QuestionSet:
        docx = _require_docx()
        document = docx.Document(source)
        assets_dir = Path(options.get("assets_dir") or source.with_name(f"{source.stem}_assets"))
        flat_questions: list[Question] = []
        parent_ids: dict[str, str] = {}
        current: Question | None = None
        current_parent_id: str | None = None
        has_marked_questions = any(
            paragraph.text.strip().startswith("Question ")
            and paragraph.style.name.startswith("Heading")
            for paragraph in document.paragraphs
        )
        for paragraph in document.paragraphs:
            value = paragraph.text.strip()
            if value.startswith("Question ") and paragraph.style.name.startswith("Heading"):
                if current:
                    flat_questions.append(current)
                    if current_parent_id:
                        parent_ids[current.id] = current_parent_id
                current = Question(
                    value[len("Question ") :].strip(), "short_answer", Content(), "en"
                )
                current_parent_id = None
                continue
            if current is None and (
                not value
                or has_marked_questions
                or paragraph.style.name in {"Title", "Subtitle"}
                or value == (document.core_properties.title or "")
            ):
                continue
            if current is None and value:
                current = Question(f"q{len(flat_questions) + 1}", "short_answer", Content(), "en")
            if current is None:
                continue
            if value.startswith("Type: "):
                fields = dict(
                    part.strip().split(": ", 1) for part in value.split("|") if ": " in part
                )
                current.type = fields.get("Type", current.type)
                current.language = fields.get("Language", current.language)
                current_parent_id = fields.get("Parent")
                continue
            if value.startswith("Question: "):
                current.stem = self._paragraph_content(
                    paragraph, value[len("Question: ") :], current, assets_dir
                )
            elif value.startswith("[") and "]" in value:
                choice_id, choice_text = value[1:].split("]", 1)
                current.choices.append(Choice(choice_id.strip(), Content.text(choice_text.strip())))
            elif value.startswith("Answer: "):
                raw = value[len("Answer: ") :]
                kind, separator, answer_value = raw.partition(":")
                if separator:
                    parsed_value: Any = (
                        [
                            part.strip()
                            for part in answer_value.strip("[]").replace("'", "").split(",")
                        ]
                        if kind == "choices"
                        else answer_value
                    )
                    current.answer = Answer(kind, parsed_value)
                else:
                    current.answer = Answer("text", raw)
            elif value.startswith("Solution: "):
                current.solution = self._paragraph_content(
                    paragraph, value[len("Solution: ") :], current, assets_dir
                )
            elif value.startswith("Hint: "):
                current.hints.append(
                    self._paragraph_content(paragraph, value[len("Hint: ") :], current, assets_dir)
                )
            elif value and not current.stem.blocks:
                current.stem = Content.text(value)
        if current:
            flat_questions.append(current)
            if current_parent_id:
                parent_ids[current.id] = current_parent_id
        by_id = {question.id: question for question in flat_questions}
        questions: list[Question] = []
        for question in flat_questions:
            parent = by_id.get(parent_ids.get(question.id, ""))
            if parent is None:
                questions.append(question)
            else:
                parent.subquestions.append(question)
        return QuestionSet(
            id=source.stem,
            title=document.core_properties.title or source.stem,
            questions=questions,
            language="en",
        )

    def _paragraph_content(
        self, paragraph, text: str, question: Question, assets_dir: Path
    ) -> Content:
        blocks = [ContentBlock(type="text", text=text)] if text else []
        for element in paragraph._p.iter():
            if not element.tag.endswith("}blip"):
                continue
            relationship_id = element.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if not relationship_id:
                continue
            part = paragraph.part.related_parts.get(relationship_id)
            if part is None:
                continue
            blob = part.blob
            digest = hashlib.sha256(blob).hexdigest()
            suffix = Path(str(getattr(part, "partname", "image.bin"))).suffix or ".bin"
            asset_id = f"asset-{digest[:12]}"
            assets_dir.mkdir(parents=True, exist_ok=True)
            output_path = assets_dir / f"{asset_id}{suffix}"
            output_path.write_bytes(blob)
            relative_uri = f"{assets_dir.name}/{output_path.name}"
            if all(asset.id != asset_id for asset in question.assets):
                question.assets.append(
                    Asset(
                        asset_id,
                        "image",
                        relative_uri,
                        mimetypes.guess_type(output_path.name)[0],
                        digest,
                    )
                )
            blocks.append(ContentBlock(type="image", asset_id=asset_id, alt_text="Imported image"))
        return Content(blocks)
