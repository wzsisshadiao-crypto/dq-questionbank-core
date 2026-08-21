"""Build deterministic, synthetic import-case bundles shipped with the package."""

from __future__ import annotations

import hashlib
import json
import shutil
import tempfile
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import parse_xml
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "src" / "dq_questionbank" / "data" / "import_cases"
FIXED_ZIP_TIME = (2020, 1, 1, 0, 0, 0)


def canonical_bytes(payload: Any) -> bytes:
    return json.dumps(
        payload, ensure_ascii=False, sort_keys=True, separators=(",", ":")
    ).encode("utf-8")


def digest_payload(payload: Any) -> str:
    return hashlib.sha256(canonical_bytes(payload)).hexdigest()


def write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
        newline="\n",
    )


def normalize_docx(path: Path) -> None:
    with tempfile.TemporaryDirectory() as temporary:
        normalized = Path(temporary) / path.name
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
            normalized, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9
        ) as target:
            for name in sorted(source.namelist()):
                info = zipfile.ZipInfo(name, FIXED_ZIP_TIME)
                info.compress_type = zipfile.ZIP_DEFLATED
                info.external_attr = 0o600 << 16
                target.writestr(info, source.read(name))
        shutil.copyfile(normalized, path)


def write_docx(path: Path, title: str, lines: list[str], *, omml: bool = False) -> None:
    document = Document()
    document.core_properties.title = title
    document.core_properties.author = "DQ QuestionBank Core contributors"
    document.core_properties.created = document.core_properties.modified
    document.add_heading(title, level=1)
    for line in lines:
        document.add_paragraph(line)
    if omml:
        paragraph = document.add_paragraph("Native equation: ")
        equation = parse_xml(
            '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<m:r><m:t>x</m:t></m:r><m:r><m:t>^2</m:t></m:r>'
            '<m:r><m:t>=9</m:t></m:r></m:oMath>'
        )
        paragraph._p.append(equation)
    document.save(path)
    normalize_docx(path)


def write_pdf(path: Path) -> None:
    styles = getSampleStyleSheet()
    document = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title="Synthetic PDF intake worksheet",
        author="DQ QuestionBank Core contributors",
        invariant=1,
    )
    story = [
        Paragraph("Synthetic PDF Intake Worksheet", styles["Title"]),
        Spacer(1, 8 * mm),
        Paragraph("Question P-01. If 3x + 2 = 14, what is x?", styles["BodyText"]),
        Spacer(1, 4 * mm),
        Table(
            [["A", "2"], ["B", "3"], ["C", "4"], ["D", "6"]],
            colWidths=[18 * mm, 42 * mm],
            style=TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8EEF4")),
                    ("ALIGN", (0, 0), (0, -1), "CENTER"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 5),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ]
            ),
        ),
        Spacer(1, 6 * mm),
        Paragraph("Answer key: C. Evidence locator: page 1, question P-01.", styles["BodyText"]),
    ]
    document.build(story)


def content(text: str) -> dict[str, Any]:
    return {"blocks": [{"type": "text", "text": text}]}


def question_set(case_id: str, title: str, questions: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "schema_version": "1.0",
        "id": f"case-{case_id}",
        "title": title,
        "language": "en",
        "questions": questions,
    }


def evidence(question_id: str, field: str, source: str, locator: str, excerpt: str) -> dict:
    return {
        "question_id": question_id,
        "field": field,
        "source_path": source,
        "locator": locator,
        "excerpt": excerpt,
        "excerpt_sha256": hashlib.sha256(excerpt.encode("utf-8")).hexdigest(),
    }


def reference(case_dir: Path, name: str) -> dict[str, str]:
    return {"path": name, "sha256": hashlib.sha256((case_dir / name).read_bytes()).hexdigest()}


def build_case(spec: dict[str, Any]) -> None:
    case_dir = OUTPUT / spec["id"]
    case_dir.mkdir(parents=True, exist_ok=True)
    source_name = spec["source_name"]
    source_path = case_dir / source_name
    if source_path.suffix == ".json":
        write_json(source_path, spec["source_payload"])
    elif source_path.suffix == ".pdf":
        write_pdf(source_path)
    else:
        write_docx(
            source_path,
            spec["title"],
            spec["source_lines"],
            omml=bool(spec.get("omml")),
        )

    write_json(case_dir / "records.json", {"records": spec["records"]})
    write_json(case_dir / "evidence.json", {"evidence": spec["evidence"]})

    base_questions = []
    for record in spec["records"]:
        item = {
            "schema_version": "1.0",
            "id": record["external_id"],
            "type": spec["defaults"]["type"],
            "language": "en",
            "subject": record["subject"],
            "stem": content(record["prompt"]),
            "answer": {"kind": spec["answer_kind"], "value": record["answer"]},
            "solution": content(record["solution"]),
            "tags": list(record["tags"]),
        }
        if "choices" in record:
            item["choices"] = [
                {"id": choice["id"], "content": content(choice["text"])}
                for choice in record["choices"]
            ]
        base_questions.append(item)
    base = question_set(spec["id"], spec["title"], base_questions)

    proposal_ref = None
    candidate_questions = json.loads(json.dumps(base_questions))
    changes = spec.get("changes", [])
    if changes:
        proposal = {"base_sha256": digest_payload(base), "changes": changes}
        write_json(case_dir / "proposal.json", proposal)
        proposal_ref = reference(case_dir, "proposal.json")
        by_id = {item["id"]: item for item in candidate_questions}
        for change in changes:
            by_id[change["question_id"]][change["field"]] = change["after"]

    decisions = {"decisions": []}
    for item in candidate_questions:
        decision = spec.get("decision_by_id", {}).get(item["id"], "accepted")
        decisions["decisions"].append(
            {
                "question_id": item["id"],
                "candidate_sha256": digest_payload(item),
                "decision": decision,
                "note": "Synthetic case review completed.",
            }
        )
    write_json(case_dir / "decisions.json", decisions)

    mapping = {
        "id": {"path": "external_id", "required": True},
        "subject": {"path": "subject", "required": True},
        "stem": {"path": "prompt", "transform": "content", "required": True},
        "answer": {"path": "answer", "transform": spec["answer_transform"], "required": True},
        "solution": {"path": "solution", "transform": "content", "required": True},
        "tags": {"path": "tags", "transform": "string_list"},
    }
    if any("choices" in record for record in spec["records"]):
        mapping["choices"] = {"path": "choices", "transform": "choices", "required": True}
    manifest = {
        "bundle_version": "1.0",
        "id": spec["id"],
        "title": spec["title"],
        "route": spec["route"],
        "source": reference(case_dir, source_name),
        "records": reference(case_dir, "records.json"),
        "evidence": reference(case_dir, "evidence.json"),
        "decisions": reference(case_dir, "decisions.json"),
        "mapping": mapping,
        "defaults": spec["defaults"],
        "ignore_paths": ["capture_note", "extraction_profile"],
        "required_evidence_fields": ["stem", "answer"],
        "allowed_proposal_fields": ["subject", "tags", "difficulty", "solution"],
        "question_set": {
            "id": f"case-{spec['id']}",
            "title": spec["title"],
            "language": "en",
        },
    }
    if proposal_ref:
        manifest["proposal"] = proposal_ref
    write_json(case_dir / "bundle.json", manifest)


def main() -> int:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    specs = case_specs()
    for spec in specs:
        build_case(spec)
    write_json(
        OUTPUT / "index.json",
        [
            {
                "id": spec["id"],
                "title": spec["title"],
                "route": spec["route"],
                "source_type": Path(spec["source_name"]).suffix.removeprefix("."),
                "summary": spec["summary"],
                "has_ai_proposal": bool(spec.get("changes")),
            }
            for spec in specs
        ],
    )
    print(f"Built {len(specs)} import cases under {OUTPUT.relative_to(ROOT)}")
    return 0


def case_specs() -> list[dict[str, Any]]:
    choice_records = {
        "manual-web": {
            "external_id": "manual-001",
            "subject": "Mathematics",
            "prompt": "What is 7 + 5?",
            "choices": [
                {"id": "A", "text": "10"},
                {"id": "B", "text": "11"},
                {"id": "C", "text": "12"},
                {"id": "D", "text": "13"},
            ],
            "answer": "C",
            "solution": "Adding 7 and 5 gives 12.",
            "tags": ["arithmetic"],
            "capture_note": "Entered through a synthetic browser form.",
        },
        "web-ai": {
            "external_id": "web-ai-001",
            "subject": "Mathematics",
            "prompt": "Which number is prime?",
            "choices": [
                {"id": "A", "text": "9"},
                {"id": "B", "text": "11"},
                {"id": "C", "text": "15"},
                {"id": "D", "text": "21"},
            ],
            "answer": "B",
            "solution": "11 has no positive divisors other than 1 and itself.",
            "tags": ["numbers"],
            "capture_note": "Drafted in the browser before a bounded AI proposal.",
        },
        "coding-pdf": {
            "external_id": "pdf-001",
            "subject": "Algebra",
            "prompt": "If 3x + 2 = 14, what is x?",
            "choices": [
                {"id": "A", "text": "2"},
                {"id": "B", "text": "3"},
                {"id": "C", "text": "4"},
                {"id": "D", "text": "6"},
            ],
            "answer": "C",
            "solution": "Subtract 2, then divide 12 by 3 to get x = 4.",
            "tags": ["linear-equation"],
            "extraction_profile": "synthetic-pdf-v1",
        },
    }
    return [
        {
            "id": "manual-web",
            "title": "Manual browser intake",
            "route": "manual_web",
            "source_name": "form-submission.json",
            "source_payload": choice_records["manual-web"],
            "records": [choice_records["manual-web"]],
            "defaults": {"type": "single_choice", "language": "en"},
            "answer_kind": "choice",
            "answer_transform": "choice_answer",
            "summary": "A human-entered form passes through the same evidence and review boundary.",
            "evidence": [
                evidence(
                    "manual-001",
                    "stem",
                    "form-submission.json",
                    "field: prompt",
                    "What is 7 + 5?",
                ),
                evidence("manual-001", "answer", "form-submission.json", "field: answer", "C"),
            ],
        },
        {
            "id": "web-ai",
            "title": "Browser intake with bounded AI proposal",
            "route": "web_ai",
            "source_name": "browser-draft.docx",
            "source_lines": [
                "Question W-01. Which number is prime?",
                "Options: 9, 11, 15, 21",
                "Answer: 11",
            ],
            "records": [choice_records["web-ai"]],
            "defaults": {"type": "single_choice", "language": "en"},
            "answer_kind": "choice",
            "answer_transform": "choice_answer",
            "summary": "A browser draft receives only digest-bound, field-allowlisted AI suggestions.",
            "evidence": [
                evidence(
                    "web-ai-001",
                    "stem",
                    "browser-draft.docx",
                    "paragraph 2",
                    "Which number is prime?",
                ),
                evidence("web-ai-001", "answer", "browser-draft.docx", "paragraph 4", "11 (choice B)"),
            ],
            "changes": [
                {
                    "question_id": "web-ai-001",
                    "field": "tags",
                    "before": ["numbers"],
                    "after": ["numbers", "prime-numbers"],
                    "reason": "Add a narrower, reviewable taxonomy hint.",
                }
            ],
        },
        {
            "id": "coding-word",
            "title": "AI Coding intake from Word",
            "route": "ai_coding",
            "source_name": "coding-source.docx",
            "source_lines": ["Question C-01. Solve x^2 = 16 for real x.", "Answer: x = -4 or x = 4."],
            "records": [
                {
                    "external_id": "coding-001",
                    "subject": "Algebra",
                    "prompt": "Solve x^2 = 16 for real x.",
                    "answer": "x = -4 or x = 4",
                    "solution": "Taking square roots gives two real values: -4 and 4.",
                    "tags": ["equations"],
                    "extraction_profile": "synthetic-word-v1",
                }
            ],
            "defaults": {"type": "short_answer", "language": "en"},
            "answer_kind": "text",
            "answer_transform": "text_answer",
            "summary": (
                "Coding-assisted Word extraction uses a declarative map and "
                "human-reviewed proposal."
            ),
            "evidence": [
                evidence(
                    "coding-001",
                    "stem",
                    "coding-source.docx",
                    "paragraph 2",
                    "Solve x^2 = 16 for real x.",
                ),
                evidence("coding-001", "answer", "coding-source.docx", "paragraph 3", "x = -4 or x = 4"),
            ],
            "changes": [
                {
                    "question_id": "coding-001",
                    "field": "difficulty",
                    "before": None,
                    "after": 0.35,
                    "reason": "Provide a bounded estimate for human review.",
                }
            ],
        },
        {
            "id": "coding-pdf",
            "title": "AI Coding intake from PDF",
            "route": "ai_coding_pdf",
            "source_name": "worksheet.pdf",
            "records": [choice_records["coding-pdf"]],
            "defaults": {"type": "single_choice", "language": "en"},
            "answer_kind": "choice",
            "answer_transform": "choice_answer",
            "summary": "A PDF worksheet keeps page evidence while extracted records enter review.",
            "evidence": [
                evidence(
                    "pdf-001",
                    "stem",
                    "worksheet.pdf",
                    "page 1, question P-01",
                    "If 3x + 2 = 14, what is x?",
                ),
                evidence("pdf-001", "answer", "worksheet.pdf", "page 1, answer key", "C"),
            ],
            "changes": [
                {
                    "question_id": "pdf-001",
                    "field": "difficulty",
                    "before": None,
                    "after": 0.25,
                    "reason": "Flag an estimated difficulty without changing source-derived content.",
                }
            ],
        },
        {
            "id": "coding-exam-omml",
            "title": "Exam-specific AI Coding intake with OMML",
            "route": "ai_coding_exam_omml",
            "source_name": "synthetic-exam.docx",
            "source_lines": [
                "Synthetic graduate entrance practice sheet.",
                "Question E-01. Solve x^2 = 9 over the real numbers.",
                "Question E-02. Draft item intentionally rejected during review.",
            ],
            "omml": True,
            "records": [
                {
                    "external_id": "exam-001",
                    "subject": "Mathematics",
                    "prompt": "Solve x^2 = 9 over the real numbers.",
                    "answer": "x = -3 or x = 3",
                    "solution": "Factor as (x - 3)(x + 3) = 0.",
                    "tags": ["exam", "omml"],
                    "extraction_profile": "synthetic-exam-omml-v1",
                },
                {
                    "external_id": "exam-002",
                    "subject": "Mathematics",
                    "prompt": "Draft item reserved to demonstrate rejection.",
                    "answer": "not published",
                    "solution": "This candidate is rejected by the bundled reviewer.",
                    "tags": ["exam", "review-demo"],
                    "extraction_profile": "synthetic-exam-omml-v1",
                },
            ],
            "defaults": {"type": "short_answer", "language": "en"},
            "answer_kind": "text",
            "answer_transform": "text_answer",
            "summary": (
                "An exam profile preserves native OMML evidence and demonstrates "
                "candidate rejection."
            ),
            "evidence": [
                evidence(
                    "exam-001",
                    "stem",
                    "synthetic-exam.docx",
                    "paragraph 3 + OMML",
                    "Solve x^2 = 9 over the real numbers.",
                ),
                evidence(
                    "exam-001",
                    "answer",
                    "synthetic-exam.docx",
                    "native OMML equation",
                    "x^2 = 9; roots -3 and 3",
                ),
                evidence(
                    "exam-002",
                    "stem",
                    "synthetic-exam.docx",
                    "paragraph 4",
                    "Draft item intentionally rejected during review.",
                ),
                evidence("exam-002", "answer", "synthetic-exam.docx", "review fixture", "not published"),
            ],
            "changes": [
                {
                    "question_id": "exam-001",
                    "field": "difficulty",
                    "before": None,
                    "after": 0.4,
                    "reason": "Attach an exam-profile estimate for explicit review.",
                }
            ],
            "decision_by_id": {"exam-002": "rejected"},
        },
    ]


if __name__ == "__main__":
    raise SystemExit(main())
